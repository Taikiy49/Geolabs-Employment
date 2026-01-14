# services.py
from __future__ import annotations

import io
import json
import os
import re
import smtplib
from email.message import EmailMessage
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import docx
from pypdf import PdfReader
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    PageBreak,
)

import google.generativeai as genai
from google.api_core.exceptions import PermissionDenied

# -------------------------------------------------------
# Env loading + config
# -------------------------------------------------------
def load_env_and_config() -> Dict[str, Any]:
    try:
        from dotenv import load_dotenv
        env_path = Path(__file__).resolve().parent / ".env"
        load_dotenv(dotenv_path=env_path)
    except Exception:
        print("⚠️ python-dotenv not installed or .env not found; using process env only.")

    gemini_key = os.getenv("GEMINI_API_KEY", "")
    gemini_model = os.getenv("GEMINI_MODEL", "gemini-2.5-pro")

    if gemini_key:
        genai.configure(api_key=gemini_key)
    else:
        print("⚠️ GEMINI_API_KEY not set – resume autofill will use regex fallback.")

    return {
        "FLASK_SECRET_KEY": os.getenv("FLASK_SECRET_KEY", "CHANGE_ME_IN_PRODUCTION"),
        "GEMINI_API_KEY": gemini_key,
        "GEMINI_MODEL": gemini_model,
        "APPLICATION_MAIL_TO": os.getenv("APPLICATION_MAIL_TO", "tyamashita@geolabs.net"),
        "APPLICATION_MAIL_FROM": os.getenv("APPLICATION_MAIL_FROM", "no-reply@geolabs.net"),
        "SMTP_HOST": os.getenv("SMTP_HOST", ""),
        "SMTP_PORT": int(os.getenv("SMTP_PORT", "587")),
        "SMTP_USER": os.getenv("SMTP_USER", ""),
        "SMTP_PASS": os.getenv("SMTP_PASS", ""),
        "SMTP_USE_TLS": os.getenv("SMTP_USE_TLS", "true").lower() in {"1", "true", "yes"},
        "CORS_ORIGINS": [
            "https://www.geolabs-employment.net",
            "https://geolabs-employment.net",
            "http://localhost:5173",
            "http://127.0.0.1:5173",
        ],
    }

# -------------------------------------------------------
# Constants
# -------------------------------------------------------
ALLOWED_EXTENSIONS = {".pdf", ".doc", ".docx", ".txt"}

# -------------------------------------------------------
# File helpers
# -------------------------------------------------------
def get_extension(filename: str) -> str:
    filename = filename or ""
    if "." not in filename:
        return ""
    return "." + filename.rsplit(".", 1)[-1].lower()

def extract_text_from_pdf(file_storage) -> str:
    reader = PdfReader(file_storage.stream)
    pages: List[str] = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n\n".join(pages)

def extract_text_from_docx(file_storage) -> str:
    file_bytes = file_storage.read()
    stream = io.BytesIO(file_bytes)
    document = docx.Document(stream)
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)

def extract_text_from_txt(file_storage) -> str:
    data = file_storage.read()
    try:
        return data.decode("utf-8", errors="ignore")
    except Exception:
        return data.decode("latin-1", errors="ignore")

def extract_text_from_file(file_storage) -> str:
    ext = get_extension(file_storage.filename)
    if ext == ".pdf":
        return extract_text_from_pdf(file_storage)
    if ext in {".doc", ".docx"}:
        return extract_text_from_docx(file_storage)
    if ext == ".txt":
        return extract_text_from_txt(file_storage)
    raise ValueError(f"Unsupported file type: {ext or 'unknown'}")

# -------------------------------------------------------
# Minimal resume parsing kept (unchanged from your current version)
# -------------------------------------------------------
_EMAIL_RE = re.compile(r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", re.I)
_PHONE_RE = re.compile(
    r"(?:(?:\+?1[\s\-\.])?\(?\d{3}\)?[\s\-\.]?\d{3}[\s\-\.]?\d{4})(?:\s*(?:x|ext\.?)\s*\d+)?",
    re.I,
)

def parse_resume_file(file_storage, cfg: Dict[str, Any]) -> Dict[str, Any]:
    ext = get_extension(file_storage.filename)
    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {ext}")

    text = extract_text_from_file(file_storage)
    if not text.strip():
        raise ValueError("Could not extract text from resume.")

    # Keep this endpoint behavior stable for now
    email = _EMAIL_RE.search(text or "")
    phone = _PHONE_RE.search(text or "")
    parsed = {
        "contact": {
            "email": email.group(0) if email else None,
            "phone": phone.group(0) if phone else None,
        }
    }

    return {"parsed": parsed, "meta": {"filename": file_storage.filename, "mode": "simple"}}

# -------------------------------------------------------
# PDF styling (professional “form” look)
# -------------------------------------------------------
def _safe(v: Any) -> str:
    if v is None:
        return ""
    if isinstance(v, bool):
        return "Yes" if v else "No"
    s = str(v).strip()
    return s

def _styles():
    styles = getSampleStyleSheet()

    styles.add(ParagraphStyle(
        name="DocTitle",
        parent=styles["Heading1"],
        fontSize=18,
        leading=22,
        spaceAfter=8,
    ))
    styles.add(ParagraphStyle(
        name="DocSub",
        parent=styles["BodyText"],
        fontSize=10,
        leading=13,
        textColor=colors.HexColor("#475569"),
        spaceAfter=10,
    ))
    styles.add(ParagraphStyle(
        name="SectionH",
        parent=styles["Heading2"],
        fontSize=12,
        leading=15,
        spaceBefore=10,
        spaceAfter=6,
        textColor=colors.HexColor("#0f172a"),
    ))
    styles.add(ParagraphStyle(
        name="Label",
        parent=styles["BodyText"],
        fontSize=9,
        leading=11,
        textColor=colors.HexColor("#475569"),
    ))
    styles.add(ParagraphStyle(
        name="Value",
        parent=styles["BodyText"],
        fontSize=10,
        leading=13,
        textColor=colors.HexColor("#0f172a"),
    ))
    styles.add(ParagraphStyle(
        name="Fine",
        parent=styles["BodyText"],
        fontSize=8.5,
        leading=11,
        textColor=colors.HexColor("#475569"),
    ))
    styles.add(ParagraphStyle(
        name="Legal",
        parent=styles["BodyText"],
        fontSize=9.5,
        leading=12.5,
        textColor=colors.HexColor("#0f172a"),
    ))
    styles.add(ParagraphStyle(
        name="Warning",
        parent=styles["BodyText"],
        fontSize=9.5,
        leading=12.5,
        textColor=colors.HexColor("#7c2d12"),
    ))
    return styles

def _kv_block(rows: List[Tuple[str, Any]], styles, cols=(2.0*inch, 4.8*inch)) -> Table:
    """
    Professional “form line” style:
    - Left column: label
    - Right column: value with bottom rule
    """
    data = []
    for label, value in rows:
        v = _safe(value) or "—"
        data.append([
            Paragraph(f"<b>{label}</b>", styles["Label"]),
            Paragraph(v.replace("\n", "<br/>"), styles["Value"]),
        ])

    t = Table(data, colWidths=list(cols))
    t.setStyle(TableStyle([
        ("VALIGN", (0,0), (-1,-1), "TOP"),
        ("LINEBELOW", (1,0), (1,-1), 0.6, colors.HexColor("#e5e7eb")),
        ("BOTTOMPADDING", (0,0), (-1,-1), 8),
        ("TOPPADDING", (0,0), (-1,-1), 2),
    ]))
    return t

def _card_box(flowables: List[Any]) -> Table:
    """
    Wrap a section in a subtle bordered “card” so it looks like a real form.
    """
    t = Table([[flowables]], colWidths=[6.8*inch])
    t.setStyle(TableStyle([
        ("BOX", (0,0), (-1,-1), 0.8, colors.HexColor("#e5e7eb")),
        ("BACKGROUND", (0,0), (-1,-1), colors.white),
        ("LEFTPADDING", (0,0), (-1,-1), 10),
        ("RIGHTPADDING", (0,0), (-1,-1), 10),
        ("TOPPADDING", (0,0), (-1,-1), 10),
        ("BOTTOMPADDING", (0,0), (-1,-1), 10),
    ]))
    return t

def _build_doc(title: str, subtitle: str, story: List[Any]) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=LETTER,
        leftMargin=0.8*inch,
        rightMargin=0.8*inch,
        topMargin=0.7*inch,
        bottomMargin=0.7*inch,
        title=title,
        author="Geolabs, Inc.",
    )
    doc.build(story)
    return buf.getvalue()

def _header_block(styles, title: str, subtitle: str) -> List[Any]:
    return [
        Paragraph("GEOLABS, INC.", styles["DocTitle"]),
        Paragraph(title, styles["SectionH"]),
        Paragraph(subtitle, styles["DocSub"]),
        Spacer(1, 6),
    ]

def _split_paragraphs(text: str) -> List[str]:
    # Split on double newlines to keep formatting like the application screen
    text = (text or "").strip()
    if not text:
        return []
    parts = [p.strip() for p in text.split("\n\n") if p.strip()]
    return parts

# -------------------------------------------------------
# PDF builders (6 PDFs total with resume)
# -------------------------------------------------------
def build_main_application_pdf(payload: Dict[str, Any]) -> bytes:
    styles = _styles()
    form = payload.get("form") or {}
    submitted_at = payload.get("submittedAt") or ""
    tz = (payload.get("clientMeta") or {}).get("timezone") or ""

    applicant = _safe(form.get("name")) or "Applicant"
    position = _safe(form.get("position")) or "Position"

    story: List[Any] = []
    story += _header_block(
        styles,
        "Employment Application (Main Application)",
        "This PDF contains the core application fields (excludes separate EEO/Disability/Veteran/Alcohol-Drug forms).",
    )

    summary_card = [
        Paragraph("<b>Submission Summary</b>", styles["SectionH"]),
        _kv_block([
            ("Submitted At", submitted_at),
            ("Client Timezone", tz),
            ("Applicant Name", form.get("name")),
            ("Position Applying For", form.get("position")),
            ("Preferred Office Location", form.get("location")),
        ], styles),
    ]
    story.append(_card_box(summary_card))
    story.append(Spacer(1, 10))

    # Application info
    app_card = [
        Paragraph("Application Information", styles["SectionH"]),
        _kv_block([
            ("Date", form.get("date")),
            ("Position Applying For", form.get("position")),
            ("Preferred Office Location", form.get("location")),
            ("Referred By", form.get("referredBy")),
        ], styles),
    ]
    story.append(_card_box(app_card))
    story.append(Spacer(1, 10))

    # Contact info
    contact_card = [
        Paragraph("General Information", styles["SectionH"]),
        _kv_block([
            ("Full Name", form.get("name")),
            ("Email", form.get("email")),
            ("Telephone No.", form.get("phone")),
            ("Cellular No.", form.get("cell")),
            ("Address", form.get("address")),
            ("City", form.get("city")),
            ("State", form.get("state")),
            ("ZIP Code", form.get("zip")),
        ], styles),
    ]
    story.append(_card_box(contact_card))
    story.append(Spacer(1, 10))

    # Employment blocks
    employment = form.get("employment") if isinstance(form.get("employment"), list) else []
    for i, job in enumerate(employment[:3]):
        if not isinstance(job, dict):
            continue
        job_card = [
            Paragraph(f"Employment Record — Employer #{i+1}", styles["SectionH"]),
            _kv_block([
                ("Company Name / Address", job.get("company")),
                ("Phone", job.get("phone")),
                ("Position", job.get("position")),
                ("Date Employed (From)", job.get("dateFrom")),
                ("Date Employed (To)", job.get("dateTo")),
                ("Primary Duties / Responsibilities", job.get("duties")),
                ("Reason for Leaving", job.get("reasonForLeaving")),
                ("Supervisor / Title", job.get("supervisor")),
            ], styles),
        ]
        story.append(_card_box(job_card))
        story.append(Spacer(1, 10))

    # Education
    edu_card = [
        Paragraph("Education", styles["SectionH"]),
        _kv_block([
            ("Highest Level Completed", form.get("highestEducationLevel")),
            ("School Name", form.get("educationSchoolName")),
            ("School Location", form.get("educationSchoolLocation")),
            ("Degree / Program", form.get("educationDegree")),
            ("Field of Study / Emphasis", form.get("educationFieldOfStudy")),
            ("Graduation Year / Years Attended", form.get("educationYears")),
            ("Additional Education", form.get("educationAdditional")),
        ], styles),
    ]
    story.append(_card_box(edu_card))
    story.append(Spacer(1, 10))

    # Skills
    skills_card = [
        Paragraph("Skills & Qualifications", styles["SectionH"]),
        _kv_block([
            ("Years of Relevant Experience", form.get("skillsYearsExperience")),
            ("Primary Area(s) of Focus", form.get("skillsPrimaryFocus")),
            ("Technical Skills & Field / Lab Tools", form.get("skillsTechnical")),
            ("Software & Programs", form.get("skillsSoftware")),
            ("Field / Laboratory Experience", form.get("skillsFieldLab")),
            ("Communication & Team Skills", form.get("skillsCommunication")),
            ("Certifications / Training", form.get("skillsCertifications")),
        ], styles),
    ]
    story.append(_card_box(skills_card))
    story.append(Spacer(1, 10))

    # References
    refs = form.get("references") if isinstance(form.get("references"), list) else []
    ref_rows: List[Any] = [Paragraph("References", styles["SectionH"])]
    for idx, r in enumerate(refs[:3]):
        if not isinstance(r, dict):
            continue
        ref_rows.append(Paragraph(f"<b>Reference #{idx+1}</b>", styles["Label"]))
        ref_rows.append(_kv_block([
            ("Name / Title", r.get("name")),
            ("Company / Relationship", r.get("company")),
            ("Contact No.", r.get("phone")),
        ], styles))
        ref_rows.append(Spacer(1, 6))

    ref_rows.append(Paragraph("<b>Authorization to Contact References</b>", styles["Label"]))
    ref_rows.append(_kv_block([
        ("Applicant’s Initials", form.get("certifyInitials")),
    ], styles))

    story.append(_card_box(ref_rows))
    story.append(Spacer(1, 10))

    # Medical
    med_card = [
        Paragraph("Medical Information & Authorization", styles["SectionH"]),
        _kv_block([
            ("Applicant’s Initials (acknowledgment)", form.get("medInitials")),
            ("Able to perform essential functions (with/without accommodation)", form.get("ableToPerformJob")),
        ], styles),
        Paragraph(
            "Note: Applicants should not provide medical diagnoses or detailed health history in this field.",
            styles["Fine"],
        )
    ]
    story.append(_card_box(med_card))
    story.append(Spacer(1, 10))

    # Affiliations
    aff_card = [
        Paragraph("Professional Affiliations", styles["SectionH"]),
        _kv_block([
            ("Affiliations / Licenses / Memberships", form.get("affiliations")),
        ], styles),
    ]
    story.append(_card_box(aff_card))
    story.append(Spacer(1, 10))

    # Certification & Disclosures
    cert_card = [
        Paragraph("Employment Certification & Disclosures", styles["SectionH"]),
        _kv_block([
            ("FCRA Initials", form.get("fcrInitials")),
            ("Do you know anyone presently working for Geolabs?", form.get("knowEmployee")),
            ("If yes, who?", form.get("knowEmployeeName")),
            ("Application Certification Date", form.get("applicationCertificationDate")),
            ("Application Certification Signature (typed)", form.get("applicationCertificationSignature")),
        ], styles),
        Paragraph("Typed signature serves as electronic signature.", styles["Fine"]),
    ]
    story.append(_card_box(cert_card))

    return _build_doc(
        title=f"Main Application — {applicant} — {position}",
        subtitle="Main Employment Application",
        story=story,
    )

def build_eeo_pdf(payload: Dict[str, Any]) -> bytes:
    styles = _styles()
    form = payload.get("form") or {}
    legal = payload.get("legalText") or {}

    story: List[Any] = []
    story += _header_block(
        styles,
        "EEO Voluntary Self-Identification Survey (Applicant Data)",
        "Confidential — Used for EEO-1 reporting only. Voluntary; will not affect employment opportunity.",
    )

    # --- LEGAL TEXT FIRST (TOP) ---
    story.append(Paragraph("Exact Text Shown to Applicant", styles["SectionH"]))
    for p in _split_paragraphs(legal.get("eeoNotice") or ""):
        story.append(Paragraph(p, styles["Legal"]))
        story.append(Spacer(1, 6))

    story.append(Spacer(1, 14))

    # --- SIGNATURE / RESPONSES BELOW ---
    card = [
        Paragraph("Applicant Responses", styles["SectionH"]),
        _kv_block([
            ("Name (optional)", form.get("eeoName")),
            ("Date (optional)", form.get("eeoDate")),
            ("Gender (voluntary)", form.get("eeoGender")),
            ("Race / Ethnicity (voluntary)", form.get("eeoEthnicity")),
        ], styles),
    ]
    story.append(_card_box(card))
    story.append(Spacer(1, 14))

    # --- FOOTER DISCLAIMER ---
    story.append(Paragraph(
        "This document reproduces the exact text presented to the applicant during the application process. "
        "Content has not been altered, summarized, or paraphrased.",
        styles["Fine"],
    ))

    return _build_doc("EEO Voluntary Self-Identification", "", story)

def build_disability_pdf(payload: Dict[str, Any]) -> bytes:
    styles = _styles()
    form = payload.get("form") or {}
    legal = payload.get("legalText") or {}

    story: List[Any] = []
    story += _header_block(
        styles,
        "Voluntary Self-Identification of Disability (CC-305)",
        "Confidential — Federal reporting only. Voluntary; will not affect employment opportunity.",
    )

    # Legal text (TOP)
    story.append(Paragraph("Exact Text Shown to Applicant", styles["SectionH"]))
    for p in _split_paragraphs(legal.get("disabilityNotice") or ""):
        story.append(Paragraph(p, styles["Legal"]))
        story.append(Spacer(1, 6))

    story.append(Spacer(1, 14))

    # Applicant responses / signature (BELOW)
    card = [
        Paragraph("Applicant Responses", styles["SectionH"]),
        _kv_block([
            ("Name (optional)", form.get("disabilityName")),
            ("Date (optional)", form.get("disabilityDate")),
            ("Employee ID (optional)", form.get("disabilityEmployeeId")),
            ("Voluntary response", form.get("disabilityStatus")),
            ("Signature (typed)", form.get("disabilitySignature")),
            ("Signature date", form.get("disabilitySignatureDate")),
        ], styles),
        Paragraph("Typed signature serves as electronic signature.", styles["Fine"]),
    ]
    story.append(_card_box(card))
    story.append(Spacer(1, 14))

    # Footer disclaimer
    story.append(Paragraph(
        "This document reproduces the exact text presented to the applicant during the application process. "
        "Content has not been altered, summarized, or paraphrased.",
        styles["Fine"],
    ))

    return _build_doc("Disability Self-Identification (CC-305)", "", story)

def build_veteran_pdf(payload: Dict[str, Any]) -> bytes:
    styles = _styles()
    form = payload.get("form") or {}
    legal = payload.get("legalText") or {}

    story: List[Any] = []
    story += _header_block(
        styles,
        "Protected Veteran Self-Identification (VEVRAA)",
        "Confidential — Affirmative action reporting only. Voluntary; will not affect employment opportunity.",
    )

    # Legal text (TOP)
    story.append(Paragraph("Exact Text Shown to Applicant", styles["SectionH"]))
    for p in _split_paragraphs(legal.get("veteranNotice") or ""):
        story.append(Paragraph(p, styles["Legal"]))
        story.append(Spacer(1, 6))

    story.append(Spacer(1, 14))

    # Applicant responses / signature (BELOW)
    card = [
        Paragraph("Applicant Responses", styles["SectionH"]),
        _kv_block([
            ("Veteran status (voluntary)", form.get("vetStatus")),
            ("Signature (typed)", form.get("vetName")),
            ("Date", form.get("vetDate")),
        ], styles),
        Paragraph("Typed signature serves as electronic signature.", styles["Fine"]),
    ]
    story.append(_card_box(card))
    story.append(Spacer(1, 14))

    # Footer disclaimer
    story.append(Paragraph(
        "This document reproduces the exact text presented to the applicant during the application process. "
        "Content has not been altered, summarized, or paraphrased.",
        styles["Fine"],
    ))

    return _build_doc("Protected Veteran Self-Identification (VEVRAA)", "", story)

def build_alcohol_drug_pdf(payload: Dict[str, Any]) -> bytes:
    styles = _styles()
    form = payload.get("form") or {}
    legal = payload.get("legalText") or {}

    story: List[Any] = []
    story += _header_block(
        styles,
        "Agreement to Comply with Alcohol & Drug Testing Program",
        "Signed applicant agreement. Required for consideration for employment.",
    )

    # Legal text (TOP)
    story.append(Paragraph("Exact Text Shown to Applicant", styles["SectionH"]))
    for p in _split_paragraphs(legal.get("alcoholDrugProgram") or ""):
        if "ANY APPLICANT WHO IS UNWILLING" in p:
            story.append(Paragraph(p, styles["Warning"]))
        else:
            story.append(Paragraph(p, styles["Legal"]))
        story.append(Spacer(1, 6))

    story.append(Spacer(1, 14))

    # Applicant acknowledgment / signature (BELOW)
    card = [
        Paragraph("Applicant Attestation", styles["SectionH"]),
        _kv_block([
            ("Acknowledged", form.get("drugAgreementAcknowledge")),
            ("Signature (typed)", form.get("drugAgreementSignature")),
            ("Date", form.get("drugAgreementDate")),
        ], styles),
        Paragraph("Typed signature serves as electronic signature.", styles["Fine"]),
    ]
    story.append(_card_box(card))
    story.append(Spacer(1, 14))

    # Footer disclaimer
    story.append(Paragraph(
        "This document reproduces the exact text presented to the applicant during the application process. "
        "Content has not been altered, summarized, or paraphrased.",
        styles["Fine"],
    ))

    return _build_doc("Alcohol & Drug Testing Program Agreement", "", story)


def resume_to_pdf(resume_bytes: bytes, resume_filename: str) -> Tuple[bytes, str]:
    ext = get_extension(resume_filename)
    safe_base = re.sub(r"[^A-Za-z0-9._-]+", "_", Path(resume_filename).stem).strip("_") or "Resume"

    if ext == ".pdf":
        return resume_bytes, f"{safe_base}.pdf"

    # Convert doc/docx/txt to PDF with extracted text
    if ext in {".doc", ".docx"}:
        stream = io.BytesIO(resume_bytes)
        document = docx.Document(stream)
        text = "\n".join([p.text for p in document.paragraphs if p.text.strip()])
    elif ext == ".txt":
        try:
            text = resume_bytes.decode("utf-8", errors="ignore")
        except Exception:
            text = resume_bytes.decode("latin-1", errors="ignore")
    else:
        raise ValueError(f"Unsupported resume file type: {ext}")

    styles = _styles()
    story: List[Any] = []
    story += _header_block(styles, "Resume (Converted to PDF)", "Original resume was not a PDF; converted for department review.")
    for p in _split_paragraphs(text.replace("\r\n", "\n")):
        story.append(Paragraph(p.replace("\n", "<br/>"), styles["Legal"]))
        story.append(Spacer(1, 6))

    pdf = _build_doc("Resume", "", story)
    return pdf, f"{safe_base}.pdf"

# -------------------------------------------------------
# Email sending (6 PDFs total)
# -------------------------------------------------------
def send_application_email(
    payload: Dict[str, Any],
    cfg: Dict[str, Any],
    resume_bytes: Optional[bytes] = None,
    resume_filename: Optional[str] = None,
) -> None:
    if not cfg.get("SMTP_HOST"):
        raise RuntimeError("SMTP_HOST is not configured on the server.")

    form = payload.get("form") or {}
    applicant_name = form.get("name") or "Applicant"
    position = form.get("position") or "Unknown Position"
    applicant_email = form.get("email") or "No email"

    # PDFs
    main_pdf = build_main_application_pdf(payload)
    eeo_pdf = build_eeo_pdf(payload)
    disability_pdf = build_disability_pdf(payload)
    veteran_pdf = build_veteran_pdf(payload)
    drug_pdf = build_alcohol_drug_pdf(payload)

    resume_pdf_bytes = None
    if resume_bytes and resume_filename:
        resume_pdf_bytes, _ = resume_to_pdf(resume_bytes, resume_filename)

    msg = EmailMessage()
    msg["From"] = cfg["APPLICATION_MAIL_FROM"]
    msg["To"] = cfg["APPLICATION_MAIL_TO"]
    msg["Subject"] = f"New Employment Application: {applicant_name} — {position}"
    msg.set_content(
        "\n".join(
            [
                "A new employment application was submitted from the web form.",
                "",
                f"Name: {applicant_name}",
                f"Position: {position}",
                f"Applicant Email: {applicant_email}",
                "",
                "Attached PDFs:",
                "1) Main Application",
                "2) EEO (Voluntary)",
                "3) Disability (Voluntary)",
                "4) Veteran (Voluntary)",
                "5) Alcohol/Drug Agreement",
                "6) Resume (PDF, if provided)",
            ]
        )
    )

    safe_app = re.sub(r"[^A-Za-z0-9._-]+", "_", str(applicant_name)).strip("_") or "Applicant"
    safe_pos = re.sub(r"[^A-Za-z0-9._-]+", "_", str(position)).strip("_") or "Position"

    msg.add_attachment(main_pdf, maintype="application", subtype="pdf",
                       filename=f"1_Main_Application_{safe_app}_{safe_pos}.pdf")
    msg.add_attachment(eeo_pdf, maintype="application", subtype="pdf",
                       filename=f"2_EEO_{safe_app}_{safe_pos}.pdf")
    msg.add_attachment(disability_pdf, maintype="application", subtype="pdf",
                       filename=f"3_Disability_{safe_app}_{safe_pos}.pdf")
    msg.add_attachment(veteran_pdf, maintype="application", subtype="pdf",
                       filename=f"4_Veteran_{safe_app}_{safe_pos}.pdf")
    msg.add_attachment(drug_pdf, maintype="application", subtype="pdf",
                       filename=f"5_Alcohol_Drug_{safe_app}_{safe_pos}.pdf")

    if resume_pdf_bytes:
        msg.add_attachment(resume_pdf_bytes, maintype="application", subtype="pdf",
                           filename=f"6_Resume_{safe_app}_{safe_pos}.pdf")

    with smtplib.SMTP(cfg["SMTP_HOST"], cfg["SMTP_PORT"], timeout=20) as server:
        if cfg["SMTP_USE_TLS"]:
            server.starttls()
        if cfg["SMTP_USER"] and cfg["SMTP_PASS"]:
            server.login(cfg["SMTP_USER"], cfg["SMTP_PASS"])
        server.send_message(msg)

def submit_application_payload(
    payload: Dict[str, Any],
    cfg: Dict[str, Any],
    resume_bytes: Optional[bytes] = None,
    resume_filename: Optional[str] = None,
) -> None:
    if resume_filename:
        ext = get_extension(resume_filename)
        if ext not in ALLOWED_EXTENSIONS:
            raise ValueError(f"Unsupported resume file type: {ext}")

    try:
        send_application_email(payload, cfg, resume_bytes=resume_bytes, resume_filename=resume_filename)
    except smtplib.SMTPAuthenticationError:
        raise RuntimeError("SMTP authentication failed. Check SMTP_USER/SMTP_PASS or use an app password.")
